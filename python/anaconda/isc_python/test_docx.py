'''
conda install conda-forge::python-docx
'''

from docx import Document
from docx.shared import Inches, Pt, RGBColor

document = Document()

# Presentation
document.add_heading('Le quatre-quarts', level=0)

paragraphe = document.add_paragraph('Le gateau 4/4 est un dessert classique. Les ingredients sont : ')
run = paragraphe.add_run('farine, beurre, sucre et oeufs.')
run.bold = True
paragraphe.add_run("C'est delicieux.")

# Pictures
document.add_picture('test.png', width=Inches(2.0))

# Les ingredients
document.add_heading('Les ingredients', level=1)
ingredients = [
    ('Oeufs', '4'),
    ('Farine', '100 g'),
    ('Beurre', '100 g'),
    ('Sucre', '100 g')
]
table = document.add_table(rows=1, cols=2)
ligne = table.rows[0].cells
ligne[0].text = 'Ingredients'
ligne[1].text = 'Quantite'

for ing, qte in ingredients:
    new_line = table.add_row().cells
    new_line[0].text = ing
    new_line[1].text = qte

table.style = 'Colorful Grid Accent 4'

# La recette
document.add_heading('La recette', level=1)

paragraphe2 = document.add_paragraph('Prechauffer le four : ')
run = paragraphe2.add_run('180 deg Cel')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
paragraphe2.add_run("Melanger les ingredients.")

document.save('quatre_quarts.docx')